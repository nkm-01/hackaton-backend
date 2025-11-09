"""Сервисный слой для работы с документами"""
import os
from typing import Optional
from django.core.files.uploadedfile import UploadedFile

# Импорты для работы с разными форматами документов
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

from .models import Document
from integrations.load_documents import get_document_processor


class DocumentService:
    """Сервис для работы с документами"""
    
    def __init__(self):
        self.document_processor = get_document_processor()
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF"""
        if not PyPDF2:
            raise ImportError("PyPDF2 не установлен")
        
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        
        return '\n'.join(text)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX"""
        if not DocxDocument:
            raise ImportError("python-docx не установлен")
        
        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        return '\n'.join(text)
    
    def extract_text_from_rtf(self, file_path: str) -> str:
        """Извлечение текста из RTF"""
        if not rtf_to_text:
            raise ImportError("striprtf не установлен")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            rtf_content = file.read()
            return rtf_to_text(rtf_content)
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Извлечение текста из TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def extract_text(self, document: Document) -> str:
        """
        Извлечение текста из документа в зависимости от типа
        
        Args:
            document: Объект документа
            
        Returns:
            str: Извлеченный текст
        """
        file_path = document.file.path
        
        if document.file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif document.file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif document.file_type == 'rtf':
            return self.extract_text_from_rtf(file_path)
        elif document.file_type == 'txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {document.file_type}")
    
    def get_pages_count(self, document: Document) -> Optional[int]:
        """
        Получить количество страниц документа (если применимо)
        
        Args:
            document: Объект документа
            
        Returns:
            int или None: Количество страниц
        """
        if document.file_type == 'pdf' and PyPDF2:
            try:
                with open(document.file.path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            except Exception:
                return None
        elif document.file_type == 'docx' and DocxDocument:
            try:
                # DOCX не имеет явного количества страниц
                # Можно приблизительно оценить по параграфам
                doc = DocxDocument(document.file.path)
                # Примерно 30 параграфов на страницу (грубая оценка)
                return max(1, len(doc.paragraphs) // 30)
            except Exception:
                return None
        elif document.file_type == 'txt':
            try:
                # Оценка страниц для TXT по количеству символов
                # Примерно 2000 символов на страницу
                with open(document.file.path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                    return max(1, len(text) // 2000)
            except Exception:
                return None
        
        return None
    
    def process_document(self, document: Document):
        """
        Обработка и индексация документа
        
        Args:
            document: Объект документа для обработки
        """
        try:
            # Обновление статуса
            document.status = 'processing'
            document.save()
            
            # Получение количества страниц
            pages_count = self.get_pages_count(document)
            if pages_count:
                document.pages_count = pages_count
                document.save()
            
            # Извлечение текста
            print(f"Extracting text from document {document.id}...")
            text = self.extract_text(document)
            
            # Обработка через AI модуль
            print(f"Processing document {document.id} with AI module...")
            sections = self.document_processor.process_document(text)
            
            # Индексация в Qdrant
            print(f"Indexing document {document.id} in Qdrant...")
            self.document_processor.index_document(sections, str(document.id))
            
            # Обновление статуса
            document.status = 'processed'
            document.error_message = ''
            document.save()
            
            print(f"Document {document.id} processed successfully!")
            
        except Exception as e:
            # Обработка ошибок
            print(f"Error processing document {document.id}: {str(e)}")
            document.status = 'error'
            document.error_message = str(e)
            document.save()
            raise
    
    def delete_document(self, document: Document):
        """
        Удаление документа и его данных из Qdrant
        
        Args:
            document: Объект документа для удаления
        """
        try:
            # Удаление из Qdrant
            print(f"Removing document {document.id} from Qdrant...")
            self.document_processor.remove_document(str(document.id))
            
            # Удаление файла
            if document.file and os.path.exists(document.file.path):
                os.remove(document.file.path)
            
            # Удаление записи из БД
            document.delete()
            
            print(f"Document {document.id} deleted successfully!")
            
        except Exception as e:
            print(f"Error deleting document {document.id}: {str(e)}")
            raise
    
    def reindex_document(self, document: Document):
        """
        Переиндексация документа (удаление старых данных и повторная обработка)
        
        Args:
            document: Объект документа для переиндексации
        """
        try:
            # Удаление из Qdrant
            print(f"Removing old data for document {document.id} from Qdrant...")
            self.document_processor.remove_document(str(document.id))
            
            # Повторная обработка
            self.process_document(document)
            
        except Exception as e:
            print(f"Error reindexing document {document.id}: {str(e)}")
            raise


def get_document_service() -> DocumentService:
    """Получить экземпляр сервиса документов"""
    return DocumentService()
